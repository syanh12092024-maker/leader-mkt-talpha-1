#!/usr/bin/env python3
"""Convert the TAlpha markdown document to a professional DOCX file."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x2F, 0x5C)
    return h

def add_para(text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(3)
    return p

# ============ TITLE PAGE ============
doc.add_paragraph()
doc.add_paragraph()
add_para("TALPHA", bold=True, size=36, color=RGBColor(0x0A, 0x2F, 0x5C), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Hệ thống Quản lý & Tự động hoá\nMarketing Toàn diện", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para("─" * 40, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para("Cuộc thi: Vibe Coding Challenge 2026 — Level Up Agency", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Đội thi: Pi Alpha — Tiểu Alpha Middle East", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

members = [
    "Lê Thị Diệu Thuý — Giám Đốc Pi Alpha",
    "Nguyễn Thị Phương Anh",
    "Hồ Sỹ Anh",
    "Hồ Sỹ Lộc",
    "Nguyễn Thị Hồng Ly",
]
for m in members:
    add_para(m, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

doc.add_paragraph()
add_para("Link demo: https://talpha-dashboard.vercel.app/talpha", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x22, 0x8B, 0x22))
doc.add_page_break()

# ============ SECTION 1: BỐI CẢNH ============
add_heading("1. BỐI CẢNH & VẤN ĐỀ", level=1)

add_heading("1.1. Giới thiệu đội ngũ", level=2)
add_para("Pi Alpha là đội Marketing phụ trách thị trường Middle East (UAE, KSA, Kuwait, Qatar, Bahrain, Oman) thuộc hệ thống Level Up Agency. Đội ngũ quản lý hàng chục Facebook Fanpage, phục vụ hàng ngàn khách hàng mỗi tháng với các sản phẩm chăm sóc sức khoẻ và làm đẹp. Quy trình hàng ngày bao gồm: tạo nội dung chào hàng, quản lý đơn nhập hàng từ nhà cung cấp (Trung Quốc → UAE → GCC), và gửi tin nhắn quảng bá tới khách hàng tiềm năng qua Facebook Messenger.")

add_heading("1.2. Ba nỗi đau lớn trước khi dùng AI", level=2)

add_para("Nỗi đau #1 — Không bám sát được hàng đang ở giai đoạn nào", bold=True, size=12, color=RGBColor(0xCC, 0x00, 0x00))
add_para("Quy trình nhập hàng từ Trung Quốc qua UAE rồi phân phối sang 6 nước GCC trải qua nhiều giai đoạn phức tạp. Tuy nhiên, bên mua hàng thường quên nhập liệu hoặc cập nhật trạng thái, dẫn đến:")
for b in [
    "MKT không biết hàng đã về kho đích 1 hay chưa",
    "Không phân biệt đơn chờ duyệt vs đơn đang vận chuyển",
    "Hàng đến muộn không có cảnh báo → ảnh hưởng kế hoạch bán",
    "Quản lý bằng Excel rời rạc, phải hỏi từng người → lãng phí 1–2h/ngày",
]:
    add_bullet(b)

add_para("Nỗi đau #2 — Tạo kịch bản chào hàng quá lâu, không đúng form", bold=True, size=12, color=RGBColor(0xCC, 0x00, 0x00))
add_para("Mỗi sản phẩm cần 4 đoạn kịch bản cho 4 khung giờ vàng. MKT từng dùng ChatGPT nhưng mỗi lần tạo lại ra format khác nhau, phải chỉnh 70–80% nội dung:")
for b in [
    "Mỗi page cần kịch bản riêng → phải copy-paste hướng dẫn lại từ đầu",
    "Tone không thống nhất giữa các page",
    "Tổng thời gian: 2–3 giờ chỉ để soạn cho 1 sản phẩm",
]:
    add_bullet(b)

add_para("Nỗi đau #3 — Không bắn tin cho khách ngoài 24h → thất thoát doanh thu", bold=True, size=12, color=RGBColor(0xCC, 0x00, 0x00))
add_para("Facebook giới hạn gửi tin trong 24h. Nghĩa là 70% khách tiềm năng không thể tiếp cận:")
for b in [
    "Sale dành 2h/ngày cài bot bắn tin thủ công qua Pancake CRM",
    "Bot bắn lại cho khách đã mua → spam → mất uy tín page",
    "Quên gửi vào khung giờ vàng → tỷ lệ đọc tin thấp",
]:
    add_bullet(b)

doc.add_page_break()

# ============ SECTION 2: GIẢI PHÁP ============
add_heading("2. GIẢI PHÁP — HỆ THỐNG TALPHA", level=1)
add_para("TAlpha là dashboard tích hợp 3 module, được xây dựng 100% bằng Vibe Coding — con người mô tả ý tưởng, AI viết toàn bộ code.")

add_heading("2.1. Module 1 — Tạo kịch bản chào hàng", level=2)
add_para("Cách hoạt động:", bold=True)
for b in [
    "MKT upload ảnh sản phẩm lên dashboard (kéo thả nhiều ảnh)",
    "Gemini AI API phân tích hình ảnh và sinh 4 đoạn kịch bản tự động",
    "Đoạn 1 (6h): tone lạc quan — Đoạn 2 (11h): chuyên nghiệp — Đoạn 3 (17h): FOMO — Đoạn 4 (21h): intimate",
    "MKT chỉnh sửa, lưu template, tái sử dụng cho SP tương tự",
]:
    add_bullet(b)
add_para("Kết quả: Từ 2–3 giờ → 30 giây có kịch bản hoàn chỉnh, chuẩn format 100%.", bold=True, color=RGBColor(0x00, 0x80, 0x00))

add_heading("2.2. Module 2 — Đặt và Flow hàng", level=2)
add_para("Cách hoạt động:", bold=True)
for b in [
    "Tổng quan: Dashboard real-time với số liệu đơn, cảnh báo muộn, biểu đồ quốc gia",
    "Đơn nhập hàng: Form tạo đơn đầy đủ, 17 trạng thái approval flow, nhận hàng nhiều lần (batch)",
    "Phân phối GCC: Tạo đơn phân phối cho 6 nước, tracking riêng từng nước, đối tác VC khác nhau",
]:
    add_bullet(b)
add_para("Kết quả: Thay thế hoàn toàn Excel. Từ 1–2h/ngày hỏi tiến độ → 0 phút (mở dashboard là thấy).", bold=True, color=RGBColor(0x00, 0x80, 0x00))

add_heading("2.3. Module 3 — Gửi tin nhắn hàng loạt", level=2)
add_para("Cách hoạt động:", bold=True)
for b in [
    "Soạn 4 đoạn tin nhắn cho 4 khung giờ vàng, đính kèm ảnh sản phẩm",
    "Kết nối Pancake CRM API lấy danh sách KH, lọc thông minh (loại KH đã mua)",
    "Auto Scheduler: cài 1 lần, chạy mỗi ngày. Worker 60s/lần, tự fire đúng giờ",
    "Trạng thái real-time: ✓ Xanh (đã gửi) | ⏳ Vàng (chờ) | ✗ Đỏ (lỗi)",
    "HUMAN_AGENT tag: bypass 24h Facebook lên 7 ngày qua Graph API v21",
    "Fallback thông minh: Pancake (trong 24h) → Facebook Graph API (ngoài 24h)",
]:
    add_bullet(b)
add_para("Kết quả: Sale không cần 2h/ngày cài bot. Reach 100% KH (7 ngày). Không spam KH đã mua.", bold=True, color=RGBColor(0x00, 0x80, 0x00))

doc.add_page_break()

# ============ SECTION 3: KIẾN TRÚC ============
add_heading("3. KIẾN TRÚC KỸ THUẬT", level=1)

# Tech stack table
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Layer", "Công nghệ", "Vai trò"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

data = [
    ("Frontend", "Next.js 14, React, TypeScript", "App Router, SPA dashboard"),
    ("UI/UX", "Tailwind CSS, Framer Motion", "Responsive, animation mượt"),
    ("AI", "Google Gemini API", "Sinh kịch bản từ ảnh"),
    ("CRM", "Pancake CRM REST API", "Quản lý khách hàng"),
    ("Messaging", "Facebook Graph API v21", "Gửi broadcast"),
    ("Data", "localStorage (client)", "Lưu shipment, cấu hình"),
    ("Hosting", "Vercel", "Deploy, edge network"),
]
for i, (a, b, c) in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
add_para("Kiến trúc module độc lập: mỗi module là React component riêng biệt, có state riêng, data storage riêng, không chia sẻ side-effect. Đảm bảo thêm/sửa/xoá module không ảnh hưởng module khác.", italic=True)

# ============ SECTION 4: VIBE CODING ============
add_heading("4. PHƯƠNG PHÁP — 100% VIBE CODING", level=1)
add_para("Vibe Coding là phương pháp phát triển phần mềm trong đó con người đóng vai Product Owner & Solution Architect, AI đóng vai Developer. Con người mô tả ý tưởng bằng ngôn ngữ tự nhiên; AI viết toàn bộ code, debug, và deploy.")
add_para("Quy trình:", bold=True)
for i, step in enumerate([
    "MKT mô tả pain point và quy trình hiện tại bằng tiếng Việt",
    "AI phân tích, đề xuất kiến trúc, chọn tech stack phù hợp",
    "AI viết toàn bộ component, API route, logic nghiệp vụ",
    "MKT test trên giao diện thực, feedback → AI chỉnh sửa",
    "AI cấu hình và deploy lên Vercel tự động",
], 1):
    add_bullet(f"Bước {i}: {step}")
add_para("Kết quả: 0 dòng code viết bởi con người. 100% do AI sinh ra. Phát triển < 1 tuần cho 3 module.", bold=True, color=RGBColor(0x00, 0x80, 0x00))

doc.add_page_break()

# ============ SECTION 5: KẾT QUẢ ============
add_heading("5. KẾT QUẢ & TÁC ĐỘNG", level=1)

# KPI table
kpi_table = doc.add_table(rows=8, cols=4)
kpi_table.style = 'Light Shading Accent 1'
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_headers = ["Chỉ số", "Trước TAlpha", "Sau TAlpha", "Cải thiện"]
for i, h in enumerate(kpi_headers):
    kpi_table.rows[0].cells[i].text = h
    for p in kpi_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

kpi_data = [
    ("Soạn kịch bản", "2–3 giờ/SP", "30 giây/SP", "↓ 90%"),
    ("Tần suất chào hàng", "1 lần/ngày", "4 lần/ngày", "↑ 4x"),
    ("Cài broadcast", "2 giờ/ngày", "10 phút/ngày", "↓ 85%"),
    ("Reach KH", "30% (24h)", "100% (7 ngày)", "↑ 3.3x"),
    ("Quản lý shipment", "Excel rời rạc", "Dashboard RT", "Thay thế"),
    ("Cảnh báo đơn muộn", "Không có", "Tự động", "Mới"),
    ("Code thủ công", "N/A", "0 dòng", "100% AI"),
]
for i, (a, b, c, d) in enumerate(kpi_data):
    row = kpi_table.rows[i + 1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    row.cells[3].text = d
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
add_heading("Tác động kinh doanh", level=2)
for b in [
    "Tiết kiệm ≈ 100+ giờ/tháng cho toàn đội MKT",
    "Tăng doanh thu: reach 100% KH thay vì 30%",
    "Giảm rủi ro: không spam, approval flow minh bạch",
    "Tính mở rộng: module độc lập, dễ thêm tính năng",
]:
    add_bullet(b)

add_heading("Tính sáng tạo", level=2)
for b in [
    "AI-first: AI LÀM TOÀN BỘ, không chỉ hỗ trợ",
    "Sản phẩm production thực tế, không phải demo",
    "Bypass Facebook 24h bằng HUMAN_AGENT — giải pháp nhiều agency chưa biết",
    "Non-technical team xây dựng tech product nhờ Vibe Coding",
]:
    add_bullet(b)

# ============ SECTION 6: ROADMAP ============
add_heading("6. ROADMAP PHÁT TRIỂN", level=1)

road_table = doc.add_table(rows=5, cols=3)
road_table.style = 'Light Shading Accent 1'
road_table.alignment = WD_TABLE_ALIGNMENT.CENTER
road_headers = ["Giai đoạn", "Thời gian", "Nội dung"]
for i, h in enumerate(road_headers):
    road_table.rows[0].cells[i].text = h
    for p in road_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

road_data = [
    ("V1 — Hiện tại", "Q1/2026", "3 module cốt lõi, live production"),
    ("V2 — Tối ưu", "Q2/2026", "Server scheduler, DB migration, A/B testing"),
    ("V3 — AI nâng cao", "Q3/2026", "Chatbot AI, phân tích phản hồi, đề xuất SP"),
    ("V4 — Mở rộng", "Q4/2026", "WhatsApp/Zalo, multi-brand, financial"),
]
for i, (a, b, c) in enumerate(road_data):
    row = road_table.rows[i + 1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
add_para("─" * 60, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Tài liệu được biên soạn bởi đội Pi Alpha — Tiểu Alpha Middle East, tháng 3/2026.", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
out = "/Users/macbook/Desktop/LEADER MKT TALPHA/TAlpha_Vibe_Coding_Challenge.docx"
doc.save(out)
import os
print(f"✅ Saved to: {out}")
print(f"   Size: {os.path.getsize(out) / 1024:.0f} KB")
